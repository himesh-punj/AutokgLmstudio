"""
extractors/document_loader.py
------------------------------
Loads and extracts text from:
  - PDF  (digital-native via pdfplumber)
  - PDF  (scanned / image-heavy via Surya OCR v0.17.x)
  - DOCX (via python-docx)
  - TXT  (plain read)

Also detects which pages are scanned vs text-native and identifies
pages that likely contain tables (for the table extractor).

Output per document:
    DocumentContent:
        pages: list[PageContent]   — one entry per page
        raw_text: str              — full concatenated text
        table_page_nums: list[int] — 0-indexed pages that likely have tables
        is_scanned: bool
"""

import re
import threading
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from loguru import logger
from PIL import Image

# ─── Surya OCR ────────────────────────────────────────────────────────────────
# Disabled by default — Surya v0.17.x has API instability issues.
# For digital PDFs (scanned=False), pdfplumber handles extraction perfectly.
# Re-enable by setting SURYA_ENABLED = True once you have scanned documents.
SURYA_ENABLED    = False   # ← set True to enable Surya for scanned PDFs
SURYA_AVAILABLE  = False
_surya_rec_predictor = None
_surya_det_predictor = None


# ─── Data structures ──────────────────────────────────────────────────────────

@dataclass
class PageContent:
    page_num: int       # 0-indexed physical page in the PDF
    text: str
    is_scanned: bool = False
    has_table: bool = False
    image_count: int = 0
    char_count: int = 0
    rect_count: int = 0         # number of rectangles (cell borders) on page
    has_grid_lines: bool = False # true if pdfplumber detected line-based grid
    printed_page: "int | None" = None  # page number printed in the footer/header (if detected)

    def __post_init__(self):
        self.char_count = len(self.text)


# ─── Printed page-number detection ────────────────────────────────────────────
# The physical PDF index (page_num) rarely matches the page number printed in the
# DPR (front matter, chapter dividers shift it). For human-facing reports we want
# the printed number, so we sniff it from the page header/footer.

import re as _re

_PAGE_OF   = _re.compile(r"page\s+(\d{1,4})\s+of\s+\d{1,4}", _re.I)
_PAGE_N    = _re.compile(r"\bpage\s*[:\-]?\s*(\d{1,4})\b", _re.I)
_SECTIONED = _re.compile(r"^\s*[A-Za-z]{1,3}[-.](\d{1,4})\s*$")   # e.g. "3-12", "A-5"


def _page_number_candidates(text: str) -> set[int]:
    """All plausible page-number integers in a page's header/footer zone."""
    out: set[int] = set()
    if not text:
        return out
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return out
    zone = lines[:2] + lines[-4:]
    for line in zone:
        m = _PAGE_OF.search(line) or _PAGE_N.search(line)
        if m:
            out.add(int(m.group(1)))
        if len(line) <= 16:   # short line — likely a bare footer number
            for mm in _re.finditer(r"\b(\d{1,4})\b", line):
                n = int(mm.group(1))
                if 0 < n < 3000:
                    out.add(n)
            ms = _SECTIONED.match(line)
            if ms:
                out.add(int(ms.group(1)))
    return out


def infer_page_number_offset(texts: list[str]) -> "int | None":
    """
    Find a constant offset k such that printed_page == physical_index(1-based) + k
    for as many pages as possible. A real page number increments by 1 per page, so
    the correct k is the one whose arithmetic sequence appears in the most pages'
    candidate sets — this ignores recurring header/footer noise (e.g. a repeated
    '25') that a naive per-page match would wrongly latch onto. Returns None if no
    offset explains a meaningful share of pages.
    """
    n = len(texts)
    if n == 0:
        return None
    cands = [_page_number_candidates(t) for t in texts]
    best_k, best_score = None, 0
    for k in range(-80, 11):
        score = sum(1 for i in range(n) if (i + 1 + k) in cands[i])
        if score > best_score:
            best_score, best_k = score, k
    if best_k is not None and best_score >= max(10, int(0.25 * n)):
        return best_k
    return None


def detect_printed_page(text: str) -> "int | None":
    """Single-page detection (used for DOCX/TXT). Returns the lone candidate if unambiguous."""
    cands = _page_number_candidates(text)
    return next(iter(cands)) if len(cands) == 1 else None


@dataclass
class DocumentContent:
    doc_id: str
    source_path: Path
    doc_type: str          # "pdf" | "docx" | "txt"
    pages: list[PageContent] = field(default_factory=list)
    raw_text: str = ""
    table_page_nums: list[int] = field(default_factory=list)
    is_scanned: bool = False
    total_pages: int = 0
    sector: Optional[str] = None
    sector_confidence: float = 0.0


# ─── Surya OCR singleton ──────────────────────────────────────────────────────

_surya_load_lock = threading.Lock()


def _get_surya_predictors():
    """Returns Surya predictors if enabled and available, else (None, None)."""
    global _surya_rec_predictor, _surya_det_predictor, SURYA_AVAILABLE
    if not SURYA_ENABLED:
        return None, None
    if _surya_rec_predictor is not None:
        return _surya_det_predictor, _surya_rec_predictor
    with _surya_load_lock:
        if _surya_rec_predictor is not None:
            return _surya_det_predictor, _surya_rec_predictor
        logger.info("Loading Surya OCR models...")
        try:
            from surya.recognition import RecognitionPredictor
            from surya.detection import DetectionPredictor
            _surya_det_predictor = DetectionPredictor()
            try:
                _surya_rec_predictor = RecognitionPredictor(
                    foundation_predictor=_surya_det_predictor
                )
            except TypeError:
                _surya_rec_predictor = RecognitionPredictor()
            SURYA_AVAILABLE = True
            logger.success("Surya OCR models loaded.")
        except Exception as e:
            logger.error(f"Surya model load failed: {e}")
            return None, None
    return _surya_det_predictor, _surya_rec_predictor


def _surya_ocr_page(image: Image.Image) -> str:
    """Run Surya OCR v0.17.x on a PIL Image. Returns extracted text."""
    det_predictor, rec_predictor = _get_surya_predictors()
    if det_predictor is None or rec_predictor is None:
        return ""
    try:
        # v0.17.x API: predictors are callable
        det_result = det_predictor([image])
        bboxes = [r.bboxes for r in det_result]
        rec_result = rec_predictor([image], bboxes)
        lines = []
        for page_result in rec_result:
            for line in page_result.text_lines:
                if line.text.strip():
                    lines.append(line.text)
        return "\n".join(lines)
    except Exception as e:
        logger.debug(f"Surya OCR page failed: {e}")
        return ""


# ─── Page analysis helpers ────────────────────────────────────────────────────

_TABLE_KEYWORDS = re.compile(
    r"\b(table|schedule|statement|bill of quantities|boq|summary|list of|rates?)\b",
    re.IGNORECASE
)
_MIN_TEXT_CHARS_FOR_NATIVE = 80  # fewer chars → likely scanned


def _page_likely_has_table(page: pdfplumber.page.Page, text: str) -> bool:
    """Heuristic: does this page likely contain a table?"""
    if _TABLE_KEYWORDS.search(text):
        return True
    # Check for grid lines (many horizontal/vertical lines = table)
    if len(page.lines) > 10:
        return True
    # Check for many rects (cells)
    if len(page.rects) > 6:
        return True
    return False


# ─── Per-page extraction (parallel + sequential) ──────────────────────────────

def _extract_pages_chunk(args):
    """
    Process-pool worker: extract per-page data for a list of page indices.
    Uses pdfplumber so the output is identical to the sequential path. Must stay a
    top-level function (picklable) and only reference module-level helpers.
    """
    pdf_path, indices = args
    import pdfplumber as _pp
    results = []
    with _pp.open(pdf_path) as pdf:
        n = len(pdf.pages)
        for i in indices:
            if i >= n:
                continue
            page = pdf.pages[i]
            text = page.extract_text() or ""
            results.append({
                "page_num":       i,
                "text":           text,
                "image_count":    len(page.images),
                "rect_count":     len(page.rects),
                "has_grid_lines": len(page.lines) > 8,
                "has_table":      _page_likely_has_table(page, text),
                "is_scanned":     len(text.strip()) < _MIN_TEXT_CHARS_FOR_NATIVE,
            })
            try:
                page.flush_cache()
            except Exception:
                pass
    return results


def _extract_pages_parallel(path: Path, total_pages: int) -> Optional[dict]:
    """Extract all pages across a process pool. Returns {page_num: data} or None on failure."""
    import os
    from concurrent.futures import ProcessPoolExecutor
    from config.settings import EXTRACTION_WORKERS

    n_workers = max(1, min(EXTRACTION_WORKERS, (os.cpu_count() or 2)))
    if n_workers < 2:
        return None

    indices = list(range(total_pages))
    chunks = [c for c in (indices[w::n_workers] for w in range(n_workers)) if c]
    data: dict = {}
    try:
        with ProcessPoolExecutor(max_workers=n_workers) as ex:
            for res in ex.map(_extract_pages_chunk, [(str(path), c) for c in chunks]):
                for d in res:
                    data[d["page_num"]] = d
    except Exception as e:
        logger.warning(f"Parallel page extraction failed ({e}); falling back to sequential.")
        return None
    if len(data) != total_pages:
        logger.warning("Parallel extraction incomplete; falling back to sequential.")
        return None
    logger.info(f"Extracted {total_pages} pages with {n_workers} parallel workers.")
    return data


def _extract_pages_sequential(path: Path, doc: "DocumentContent"):
    """Original single-threaded pass — also handles Surya OCR for scanned pages."""
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            native_text = page.extract_text() or ""
            is_scanned_page = len(native_text.strip()) < _MIN_TEXT_CHARS_FOR_NATIVE

            if is_scanned_page and SURYA_AVAILABLE:
                pil_img = page.to_image(resolution=200).original
                page_text = _surya_ocr_page(pil_img)
                if not page_text.strip():
                    page_text = native_text
            else:
                page_text = native_text

            pc = PageContent(
                page_num=i, text=page_text, is_scanned=is_scanned_page,
                has_table=_page_likely_has_table(page, page_text),
                image_count=len(page.images), rect_count=len(page.rects),
                has_grid_lines=len(page.lines) > 8,
            )
            doc.pages.append(pc)
            if pc.has_table:
                doc.table_page_nums.append(i)
            try:
                page.flush_cache()
            except Exception:
                pass


# ─── PDF loader ───────────────────────────────────────────────────────────────

def _load_pdf(path: Path, doc_id: str) -> DocumentContent:
    doc = DocumentContent(
        doc_id=doc_id,
        source_path=path,
        doc_type="pdf",
    )

    # Sample the first few pages to decide if the whole PDF is scanned.
    with pdfplumber.open(path) as pdf:
        doc.total_pages = len(pdf.pages)
        logger.info(f"Loading PDF: {path.name} ({doc.total_pages} pages)")
        sample_chars = sum(
            len(pdf.pages[i].extract_text() or "")
            for i in range(min(5, doc.total_pages))
        )
        doc.is_scanned = (sample_chars < _MIN_TEXT_CHARS_FOR_NATIVE * 5)
    if doc.is_scanned:
        logger.info("Document appears to be scanned — will use Surya OCR.")

    # Native PDFs: extract pages in parallel (CPU-bound pdfminer work across cores).
    # Scanned PDFs needing Surya, or any pool failure, fall back to the sequential pass.
    pages_data = None
    if doc.total_pages >= 30 and not (doc.is_scanned and SURYA_AVAILABLE):
        pages_data = _extract_pages_parallel(path, doc.total_pages)

    if pages_data:
        for i in range(doc.total_pages):
            d = pages_data[i]
            pc = PageContent(
                page_num=i, text=d["text"], is_scanned=d["is_scanned"],
                has_table=d["has_table"], image_count=d["image_count"],
                rect_count=d["rect_count"], has_grid_lines=d["has_grid_lines"],
            )
            doc.pages.append(pc)
            if d["has_table"]:
                doc.table_page_nums.append(i)
    else:
        _extract_pages_sequential(path, doc)

    # Assign printed page numbers using a document-wide constant offset (robust to
    # recurring header/footer noise). printed = physical_index(1-based) + offset.
    _offset = infer_page_number_offset([p.text for p in doc.pages])
    if _offset is not None:
        for p in doc.pages:
            p.printed_page = p.page_num + 1 + _offset
        logger.info(f"Printed page offset detected: printed = physical {_offset:+d}")

    doc.raw_text = "\n\n".join(p.text for p in doc.pages if p.text)
    logger.success(
        f"PDF loaded: {doc.total_pages} pages, "
        f"{len(doc.table_page_nums)} table pages, "
        f"scanned={doc.is_scanned}"
    )
    return doc


# ─── DOCX loader ──────────────────────────────────────────────────────────────

def _load_docx(path: Path, doc_id: str) -> DocumentContent:
    doc = DocumentContent(
        doc_id=doc_id,
        source_path=path,
        doc_type="docx",
    )
    docx = DocxDocument(path)
    full_text_parts = []
    table_count = 0

    for para in docx.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)

    for i, table in enumerate(docx.tables):
        table_count += 1
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        full_text_parts.append(f"\n[TABLE {i+1}]\n{table_text}\n[/TABLE]\n")

    full_text = "\n".join(full_text_parts)
    pc = PageContent(page_num=0, text=full_text, has_table=(table_count > 0))
    doc.pages = [pc]
    doc.raw_text = full_text
    doc.total_pages = 1
    if table_count > 0:
        doc.table_page_nums = [0]

    logger.success(f"DOCX loaded: {len(docx.paragraphs)} paragraphs, {table_count} tables")
    return doc


# ─── TXT loader ───────────────────────────────────────────────────────────────

def _load_txt(path: Path, doc_id: str) -> DocumentContent:
    text = path.read_text(encoding="utf-8", errors="replace")
    pc = PageContent(page_num=0, text=text)
    return DocumentContent(
        doc_id=doc_id,
        source_path=path,
        doc_type="txt",
        pages=[pc],
        raw_text=text,
        total_pages=1,
    )


# ─── Public API ───────────────────────────────────────────────────────────────

def load_document(path: Path, doc_id: str = None) -> DocumentContent:
    """
    Load any supported document and return a DocumentContent object.
    Automatically chooses the right loader based on file extension.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    _id = doc_id or path.stem
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path, _id)
    elif suffix in (".docx", ".doc"):
        return _load_docx(path, _id)
    elif suffix == ".txt":
        return _load_txt(path, _id)
    else:
        raise ValueError(f"Unsupported document type: {suffix}")